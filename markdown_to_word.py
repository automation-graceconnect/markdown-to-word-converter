#!/usr/bin/env python3
"""
Markdown to Word Converter
Converts markdown files to Word documents with proper formatting.
- Headings mapped to standard Word styles
- Bold and italic formatting preserved
- Tables converted to Word tables
- LaTeX equations kept in $ format and centered
"""

import re
import sys
import os
import base64
import tempfile
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.xmlchemy import BaseOxmlElement
import argparse
from tkinter import Tk, filedialog

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from latex2mathml.converter import convert as latex_to_mathml
except ImportError:
    latex_to_mathml = None

try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False


class MarkdownToWordConverter:
    def __init__(self):
        self.doc = Document()
        self.last_api_call_time = 0  # Track last mermaid API call for rate limiting

        # Set margins to 0.98" on all sides
        section = self.doc.sections[0]
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)

    def insert_equation_object(self, paragraph, latex_code, is_display=False):
        """
        Insert a LaTeX equation as a Word equation object using OMML
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create the OMML (Office Math Markup Language) structure
        # This creates a proper Word equation object

        # Create oMath element
        oMath = OxmlElement('m:oMath')

        # Create a run for the equation text
        oMathRun = OxmlElement('m:r')

        # Add run properties
        rPr = OxmlElement('m:rPr')
        sty = OxmlElement('m:sty')
        sty.set(qn('m:val'), 'p')  # Professional style
        rPr.append(sty)
        oMathRun.append(rPr)

        # Add the equation text
        t = OxmlElement('m:t')
        t.text = latex_code
        oMathRun.append(t)

        oMath.append(oMathRun)

        # For display equations, wrap in oMathPara
        if is_display:
            oMathPara = OxmlElement('m:oMathPara')
            oMathPara.append(oMath)
            paragraph._element.append(oMathPara)
        else:
            # For inline equations, add directly to paragraph
            paragraph._element.append(oMath)

    def parse_inline_formatting(self, text, paragraph):
        """
        Parse inline formatting like **bold**, *italic*, and preserve LaTeX $...$
        More robust version that avoids corruption
        """
        if not text:
            return

        # Process text character by character to handle formatting
        i = 0
        current_text = ""

        while i < len(text):
            # Check for **bold**
            if i < len(text) - 1 and text[i:i+2] == '**':
                # Add accumulated text first
                if current_text:
                    run = paragraph.add_run(current_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    current_text = ""

                # Find closing **
                end = text.find('**', i + 2)
                if end != -1:
                    bold_text = text[i+2:end]
                    if bold_text:  # Only add if not empty
                        run = paragraph.add_run(bold_text)
                        run.bold = True
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                    i = end + 2
                    continue
                else:
                    # No closing **, treat as literal
                    current_text += text[i]
                    i += 1

            # Check for *italic*
            elif text[i] == '*' and i > 0 and text[i-1] != '*':
                # Make sure it's not part of **
                if i + 1 < len(text) and text[i+1] != '*':
                    # Add accumulated text first
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        current_text = ""

                    # Find closing *
                    end = i + 1
                    while end < len(text) and text[end] != '*':
                        end += 1

                    if end < len(text):
                        italic_text = text[i+1:end]
                        if italic_text:  # Only add if not empty
                            run = paragraph.add_run(italic_text)
                            run.italic = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                        i = end + 1
                        continue
                    else:
                        # No closing *, treat as literal
                        current_text += text[i]
                        i += 1
                else:
                    current_text += text[i]
                    i += 1

            # Check for $latex$
            elif text[i] == '$':
                # Add accumulated text first
                if current_text:
                    run = paragraph.add_run(current_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    current_text = ""

                # Find closing $
                end = text.find('$', i + 1)
                if end != -1:
                    latex_code = text[i+1:end]  # Extract LaTeX without $ signs
                    if latex_code:
                        # Insert as Word equation object
                        self.insert_equation_object(paragraph, latex_code, is_display=False)
                    i = end + 1
                    continue
                else:
                    # No closing $, treat as literal
                    current_text += text[i]
                    i += 1

            else:
                current_text += text[i]
                i += 1

        # Add any remaining text
        if current_text:
            run = paragraph.add_run(current_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    def add_heading(self, text, level):
        """Add a heading with the specified level"""
        text = text.strip()
        heading = self.doc.add_heading(text, level=level)
        return heading

    def add_paragraph_with_formatting(self, text):
        """Add a paragraph with inline formatting (bold, italic, LaTeX)"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Set paragraph spacing
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.08

        self.parse_inline_formatting(text, paragraph)
        return paragraph

    def add_centered_latex(self, text):
        """Add centered LaTeX equation (display mode)"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set paragraph spacing
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = 1.08

        # Remove $ signs from the equation text
        latex_code = text.strip()
        if latex_code.startswith('$$') and latex_code.endswith('$$'):
            latex_code = latex_code[2:-2]
        elif latex_code.startswith('$') and latex_code.endswith('$'):
            latex_code = latex_code[1:-1]

        # Insert as Word equation object (display mode)
        self.insert_equation_object(paragraph, latex_code, is_display=True)
        return paragraph

    def parse_table(self, lines):
        """
        Parse markdown table and convert to Word table
        Returns the number of lines consumed
        """
        if len(lines) < 2:
            return 0

        # First line is header
        header_line = lines[0].rstrip()
        separator_line = lines[1].rstrip()

        # Check if header line starts with | and contains pipes
        if not header_line.strip().startswith('|') or '|' not in header_line:
            return 0

        # Check if it's a valid table (separator line should have |---|---|)
        if not separator_line.strip().startswith('|') or not re.match(r'^\s*\|[\s\-:|]+\|', separator_line):
            return 0

        # Parse header - remove only leading/trailing empty cells from split
        headers = [cell.strip() for cell in header_line.split('|')]
        # Remove leading empty cell (before first |) and trailing empty cell (after last |)
        if headers and headers[0] == '':
            headers = headers[1:]
        if headers and headers[-1] == '':
            headers = headers[:-1]
        num_cols = len(headers)

        if num_cols == 0:
            return 0

        # Count data rows
        data_rows = []
        row_idx = 2
        while row_idx < len(lines):
            line = lines[row_idx].rstrip()

            # Stop if empty line or no pipe
            if not line.strip() or '|' not in line:
                break

            # Stop if line doesn't start with | (not a proper table row)
            if not line.strip().startswith('|'):
                break

            row_data = [cell.strip() for cell in line.split('|')]
            # Remove leading empty cell (before first |) and trailing empty cell (after last |)
            if row_data and row_data[0] == '':
                row_data = row_data[1:]
            if row_data and row_data[-1] == '':
                row_data = row_data[:-1]

            # Stop if wrong number of columns
            if len(row_data) != num_cols:
                break

            data_rows.append(row_data)
            row_idx += 1

        # Save the total lines consumed before reusing variable
        lines_consumed = row_idx

        # Create Word table
        num_rows = len(data_rows) + 1  # +1 for header
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Add header
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            paragraph = cell.paragraphs[0]
            # Clear existing content
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add text with formatting
            self.parse_inline_formatting(header_text, paragraph)
            # Make header bold (font already set to Calibri 11pt by parse_inline_formatting)
            for run in paragraph.runs:
                run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(data_rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Font already set to Calibri 11pt by parse_inline_formatting
                self.parse_inline_formatting(cell_text, paragraph)

        return lines_consumed  # Number of lines consumed

    def is_display_latex(self, line):
        """Check if line is a display mode LaTeX equation ($$...$$)"""
        stripped = line.strip()
        return stripped.startswith('$$') and stripped.endswith('$$')

    def is_inline_latex_only(self, line):
        """Check if line contains only inline LaTeX equation ($...$)"""
        stripped = line.strip()
        # Check for centered inline equations like: $...$
        if stripped.startswith('$') and stripped.endswith('$') and not stripped.startswith('$$'):
            return True
        return False

    def parse_mermaid_block(self, lines):
        """
        Parse and render Mermaid code block using mermaid.ink API + Playwright for high quality
        Returns the number of lines consumed, or 0 if not a valid mermaid block
        """
        if not REQUESTS_AVAILABLE or not PLAYWRIGHT_AVAILABLE:
            return 0

        # Check if first line is ```mermaid
        if not lines[0].strip().lower() in ['```mermaid', '``` mermaid']:
            return 0

        # Find the closing ```
        mermaid_code = []
        line_idx = 1
        while line_idx < len(lines):
            line = lines[line_idx].rstrip()
            if line.strip() == '```':
                break
            mermaid_code.append(line)
            line_idx += 1

        if line_idx >= len(lines):
            # No closing ```, invalid block
            return 0

        # We have a valid mermaid block
        mermaid_text = '\n'.join(mermaid_code)

        try:
            # Encode mermaid code to base64
            encoded = base64.b64encode(mermaid_text.encode('utf-8')).decode('utf-8')

            # Add cooldown between API calls to be respectful to mermaid.ink
            time_since_last_call = time.time() - self.last_api_call_time
            if time_since_last_call < 1.0:  # Wait at least 1 second between calls
                time.sleep(1.0 - time_since_last_call)

            # Fetch SVG from mermaid.ink
            svg_url = f"https://mermaid.ink/svg/{encoded}"
            response = requests.get(svg_url, timeout=15)
            self.last_api_call_time = time.time()

            if response.status_code != 200:
                raise Exception(f"Failed to fetch SVG (HTTP {response.status_code})")

            svg_content = response.text

            # Use Playwright to convert SVG to high-resolution PNG
            # By embedding SVG inline, we avoid CORS restrictions
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()

                # Create HTML with inline SVG
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                </head>
                <body style="margin: 0; padding: 0;">
                    <div id="svg-container">
                        {svg_content}
                    </div>
                </body>
                </html>
                """

                page.set_content(html_content)

                # Convert SVG to PNG using canvas
                # Scale 5x for high resolution (targeting 4K quality)
                png_base64 = page.evaluate("""
                    () => {
                        const svg = document.querySelector('svg');
                        if (!svg) return null;

                        const svgRect = svg.getBoundingClientRect();
                        const scale = 20;  // 20x resolution for true 4K quality

                        const canvas = document.createElement('canvas');
                        canvas.width = svgRect.width * scale;
                        canvas.height = svgRect.height * scale;

                        const ctx = canvas.getContext('2d');
                        ctx.scale(scale, scale);

                        // Serialize SVG to XML and create data URL
                        const svgData = new XMLSerializer().serializeToString(svg);
                        const svgDataUrl = 'data:image/svg+xml;base64,' + btoa(unescape(encodeURIComponent(svgData)));

                        return new Promise((resolve, reject) => {
                            const img = new Image();
                            img.onload = () => {
                                ctx.drawImage(img, 0, 0);
                                const pngData = canvas.toDataURL('image/png').split(',')[1];
                                resolve(pngData);
                            };
                            img.onerror = (e) => {
                                reject(new Error('Failed to load SVG image'));
                            };
                            img.src = svgDataUrl;
                        });
                    }
                """)

                browser.close()

            if not png_base64:
                raise Exception("Failed to convert SVG to PNG")

            # Decode base64 PNG
            png_data = base64.b64decode(png_base64)

            # Save to temp file
            temp_png = os.path.join(tempfile.gettempdir(), f'mermaid_{hash(mermaid_text)}.png')
            with open(temp_png, 'wb') as f:
                f.write(png_data)

            # Calculate appropriate size respecting both width and height constraints
            # A4 dimensions: 8.27" x 11.69", max diagram: 75% = 6.2" x 8.77"
            max_width = 6.2
            max_height = 8.77

            # Get image dimensions to calculate aspect ratio
            if PIL_AVAILABLE:
                try:
                    img = Image.open(temp_png)
                    img_width, img_height = img.size
                    aspect_ratio = img_width / img_height

                    # Calculate width needed if constrained by height
                    width_for_height = max_height * aspect_ratio

                    # Use the smaller of the two to respect both constraints
                    final_width = min(max_width, width_for_height)
                except:
                    # If PIL fails, use default width
                    final_width = max_width
            else:
                # If PIL not available, use default width
                final_width = max_width

            # Insert the rendered PNG into the document
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(temp_png, width=Inches(final_width))

            # Add some spacing after diagram
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(8)

            # Clean up temp file
            try:
                os.remove(temp_png)
            except:
                pass

            print(f"✓ Rendered high-quality Mermaid diagram")

        except Exception as e:
            print(f"Warning: Could not render Mermaid diagram: {e}")
            # Add error message to document
            error_para = self.doc.add_paragraph()
            error_run = error_para.add_run(f"[Mermaid diagram could not be rendered - requires internet connection]")
            error_run.font.name = 'Calibri'
            error_run.font.size = Pt(11)
            error_run.italic = True

        # Return number of lines consumed (including opening and closing ```)
        return line_idx + 1

    def convert(self, markdown_file, output_file=None):
        """Convert markdown file to Word document"""
        if output_file is None:
            output_file = markdown_file.rsplit('.', 1)[0] + '.docx'

        try:
            with open(markdown_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
        except Exception as e:
            print(f"Error reading file: {e}")
            return None

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Check for images: ![alt text](image_path)
            image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line.strip())
            if image_match:
                alt_text = image_match.group(1)
                image_path = image_match.group(2)

                # Add image
                try:
                    # Make path absolute if it's relative
                    if not image_path.startswith('/'):
                        import os
                        base_dir = os.path.dirname(markdown_file)
                        image_path = os.path.join(base_dir, image_path)

                    # Add the image
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(6))  # 6 inch width

                    # Add caption if alt text exists
                    if alt_text:
                        caption_para = self.doc.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(alt_text)
                        caption_run.font.name = 'Calibri'
                        caption_run.font.size = Pt(10)
                        caption_run.italic = True

                        # Set caption spacing
                        caption_format = caption_para.paragraph_format
                        caption_format.space_before = Pt(4)
                        caption_format.space_after = Pt(8)

                except Exception as e:
                    # If image fails to load, add error message
                    error_para = self.doc.add_paragraph()
                    error_run = error_para.add_run(f"[Image not found: {image_path}]")
                    error_run.font.name = 'Calibri'
                    error_run.font.size = Pt(11)
                    error_run.italic = True
                    print(f"Warning: Could not load image: {image_path} - {e}")

                i += 1
                continue

            # Check for mermaid code blocks
            if line.strip().lower() in ['```mermaid', '``` mermaid']:
                lines_consumed = self.parse_mermaid_block(lines[i:])
                if lines_consumed > 0:
                    i += lines_consumed
                    continue

            # Check for headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.add_heading(text, level)
                i += 1
                continue

            # Check for horizontal rule (---) and skip it
            if line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
                i += 1
                continue

            # Check for display mode LaTeX ($$...$$)
            if self.is_display_latex(line):
                self.add_centered_latex(line.strip())
                i += 1
                continue

            # Check for inline LaTeX that should be centered
            if self.is_inline_latex_only(line):
                self.add_centered_latex(line.strip())
                i += 1
                continue

            # Check for tables
            if '|' in line and i + 1 < len(lines):
                lines_consumed = self.parse_table(lines[i:])
                if lines_consumed > 0:
                    i += lines_consumed
                    continue

            # Check for bullet lists
            bullet_match = re.match(r'^[\s]*[-*+]\s+(.+)$', line)
            if bullet_match:
                text = bullet_match.group(1)
                paragraph = self.doc.add_paragraph(style='List Bullet')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Set paragraph spacing and indentation
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.5)  # Total left indent
                paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(8)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph_format.line_spacing = 1.08

                self.parse_inline_formatting(text, paragraph)

                # Set font for all runs to Calibri 11pt
                for run in paragraph.runs:
                    run.font.size = Pt(11)

                i += 1
                continue

            # Check for numbered lists (including bold-formatted: **1. Text:** or plain: 1. Text)
            # Try to match **1. format first
            numbered_match_bold = re.match(r'^[\s]*\*\*\d+\.\s+(.+)$', line)
            # Then try regular 1. format
            numbered_match_plain = re.match(r'^[\s]*\d+\.\s+(.+)$', line)

            text = None
            if numbered_match_bold:
                # For **1. Text**, extract text after number and add ** back for complete bold formatting
                # The line is: **1. Measurement Error:** so we want: **Measurement Error:**
                text = '**' + numbered_match_bold.group(1)
            elif numbered_match_plain:
                # For plain 1. Text, just extract the text after the number
                text = numbered_match_plain.group(1)

            if text is not None:
                paragraph = self.doc.add_paragraph(style='List Number')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Set paragraph spacing and indentation
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Inches(0.5)  # Total left indent
                paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(8)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                paragraph_format.line_spacing = 1.08

                self.parse_inline_formatting(text, paragraph)

                # Set font for all runs to Calibri 11pt
                for run in paragraph.runs:
                    run.font.size = Pt(11)

                i += 1
                continue

            # Regular paragraph
            if line.strip():  # Only add non-empty lines
                self.add_paragraph_with_formatting(line)
            i += 1

        # Save the document
        try:
            self.doc.save(output_file)
            print(f"✓ Converted '{markdown_file}' to '{output_file}'")
            return output_file
        except Exception as e:
            print(f"Error saving document: {e}")
            return None


def main():
    input_file = None
    output_file = None

    # Check if file was provided via command line (drag & drop or terminal)
    if len(sys.argv) > 1:
        # Command-line mode
        parser = argparse.ArgumentParser(
            description='Convert Markdown to Word document with formatting'
        )
        parser.add_argument(
            'input_file',
            help='Input markdown file (.md)'
        )
        parser.add_argument(
            '-o', '--output',
            help='Output Word file (.docx). Default: same name as input with .docx extension'
        )

        args = parser.parse_args()
        input_file = args.input_file
        output_file = args.output
    else:
        # GUI mode - show file picker
        root = Tk()
        root.withdraw()  # Hide the main window
        root.attributes('-topmost', True)  # Bring dialog to front

        print("Please select a markdown file to convert...")
        input_file = filedialog.askopenfilename(
            title='Select Markdown File',
            filetypes=[
                ('Markdown files', '*.md'),
                ('All files', '*.*')
            ]
        )
        root.destroy()

        if not input_file:
            print("No file selected. Exiting.")
            input("Press Enter to close...")
            return

    # Convert the file
    converter = MarkdownToWordConverter()
    result = converter.convert(input_file, output_file)

    # If running as GUI (no command-line args), show success message and wait
    if len(sys.argv) <= 1:
        if result:
            print(f"\n✓ Success! Created: {result}")
        else:
            print("\n✗ Conversion failed. Please check the error messages above.")
        input("\nPress Enter to close...")


if __name__ == '__main__':
    main()
